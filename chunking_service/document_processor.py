import logging
import os
import sys
import time
import shutil
import math
import requests
from pathlib import Path
import re
from datetime import datetime
from langchain_text_splitters import RecursiveCharacterTextSplitter
from dotenv import load_dotenv

import fitz  # PyMuPDF
import zipfile
import io
from docx import Document

# Add project root to Python path
project_root = Path(__file__).parent.parent
sys.path.insert(0, str(project_root))

from ingest_service.documents_ingestion import DocumentLoader
from embedding_service.embedding_processor import embedding

# Load environment variables
load_dotenv(project_root / ".env")


def read_docx_from_zip(zip_file, filename: str) -> str:
    """
    从 ZIP 文件中读取 .docx 文件内容，包括文本和表格。

    Args:
        zip_file: ZipFile 对象
        filename: .docx 文件名

    Returns:
        提取的文本内容（表格转为 Markdown 格式）
    """
    # 读取 docx 文件
    doc_bytes = zip_file.read(filename)
    doc = Document(io.BytesIO(doc_bytes))

    content_parts = []

    # 遍历所有段落和表格，按顺序提取
    for element in doc.element.body:
        # 判断是段落还是表格
        if element.tag.endswith('p'):  # paragraph
            for para in doc.paragraphs:
                if para._element == element:
                    if para.text.strip():
                        content_parts.append(para.text.strip())
                    break
        elif element.tag.endswith('tbl'):  # table
            for table in doc.tables:
                if table._element == element:
                    # 将表格转换为 Markdown 格式
                    md_table = convert_table_to_markdown(table)
                    if md_table:
                        content_parts.append(md_table)
                    break

    return '\n\n'.join(content_parts)


def convert_table_to_markdown(table) -> str:
    """
    将 python-docx 的表格对象转换为 Markdown 格式。

    Args:
        table: python-docx Table 对象

    Returns:
        Markdown 格式的表格字符串
    """
    if not table.rows:
        return ''

    # 提取所有单元格文本
    rows_data = []
    for row in table.rows:
        row_text = [cell.text.strip() for cell in row.cells]
        rows_data.append(row_text)

    if not rows_data:
        return ''

    # 构建 Markdown 表格
    md_lines = []

    # 第一行作为表头
    header = rows_data[0]
    md_lines.append('| ' + ' | '.join(header) + ' |')

    # 分隔线
    separator = '| ' + ' | '.join(['---'] * len(header)) + ' |'
    md_lines.append(separator)

    # 数据行
    for row in rows_data[1:]:
        md_lines.append('| ' + ' | '.join(row) + ' |')

    return '\n'.join(md_lines)


def split_into_chunks(text: str, max_chars: int = 500, overlap_ratio: float = 0.3) -> list:
    """
    Split text into semantically meaningful chunks with overlap.
    
    Args:
        text: Input text to chunk
        max_chars: Maximum characters per chunk (approx 512 tokens for Chinese)
        overlap_ratio: Percentage of overlap between chunks
        
    Returns:
        List of text chunks
    """
    # First try to split into paragraphs (using empty lines as separators)
    paragraphs = [p.strip() for p in text.split('\n\n') if p.strip()]
    
    # If no paragraphs found, try splitting by sentence endings
    if len(paragraphs) <= 1:
        sentences = re.split(r'(?<=[。！？.!?])\s+', text)
        paragraphs = [s.strip() for s in sentences if s.strip()]
    
    chunks = []
    current_chunk = ""
    overlap_size = int(max_chars * overlap_ratio)
    
    for para in paragraphs:
        # Ensure each paragraph can fit in a chunk
        if len(para) > max_chars:
            # If paragraph is too long, split it
            words = para.split()
            temp_chunk = ""
            for word in words:
                if len(temp_chunk) + len(word) + 1 > max_chars:
                    if temp_chunk:
                        chunks.append(temp_chunk)
                        # Create overlap
                        temp_chunk = temp_chunk[-overlap_size:] + " " + word
                    else:
                        # Single word longer than max_chars
                        chunks.append(word)
                else:
                    temp_chunk += (" " + word if temp_chunk else word)
            if temp_chunk:
                current_chunk = temp_chunk
        else:
            # If adding this paragraph exceeds max length
            if len(current_chunk) + len(para) > max_chars and current_chunk:
                chunks.append(current_chunk)
                # Create overlap by carrying over part of current chunk
                current_chunk = current_chunk[-overlap_size:] + " " + para
            else:
                current_chunk += ("\n" + para if current_chunk else para)
    
    # Add any remaining content as final chunk
    if current_chunk:
        chunks.append(current_chunk)
    
    # Final check: ensure no chunk exceeds max length
    final_chunks = []
    for chunk in chunks:
        if len(chunk) > max_chars:
            # Split oversized chunk into smaller ones
            for i in range(0, len(chunk), max_chars):
                final_chunks.append(chunk[i:i+max_chars])
        else:
            final_chunks.append(chunk)
    
    return final_chunks

def split_into_chunks_v2(text: str) -> list:
    text_splitter = RecursiveCharacterTextSplitter(
        separators = ["\n\n", "\n", "。", "！", "？", ".", "!", "?", " ", ""],
        chunk_size=1024,
        chunk_overlap=200,
        length_function=len,
        is_separator_regex=False,
    )
    return text_splitter.create_documents([text])

def split_into_chunks_pdf(text: str, max_chars: int = 500, overlap_ratio: float = 0.3) -> list:
    return None
def process_documents(directory: str) -> list:
    """
    Process all documents in directory: load and chunk them.
    
    Args:
        directory: Path to directory with source documents or a specific .txt file
        
    Returns:
        List of chunked document objects with metadata
    """
    documents = DocumentLoader(directory, file_type='txt').load_documents()
    processed = []
    for doc in documents:
        chunks = split_into_chunks(doc['content'])
        for i, chunk in enumerate(chunks):
            processed.append({
                'doc_id': f"{doc['id']}_chunk_{i}",
                'content': chunk,
                'source': doc['source'],
                'original_id': doc['id'],
                'chunk_index': i,
                'total_chunks': len(chunks)
            })
    
    print(f"Processed {len(processed)} chunks from {len(documents)} documents\n")
    return processed

def process_documents_v2(directory: str) -> list:
    documents = DocumentLoader(directory, file_type='txt').load_documents()
    processed = []
    for doc in documents:
        chunks = split_into_chunks_v2(doc['content'])
        logging.debug(chunks)
        for i, chunk in enumerate(chunks):
            processed.append({
                'doc_id': f"{doc['id']}_chunk_{i}",
                'content': chunk,
                'source': doc['source'],
                'original_id': doc['id'],
                'chunk_index': i,
            })
    print(f"Processed {len(processed)} chunks from {len(documents)} documents\n")
    print(type(processed))
    return processed

def process_documents_pdf(directory: str) -> list:
    """
    Process PDF documents using MinerU API.

    Steps:
    1. Get PDF file list from directory
    2. For each PDF:
       - Check page count
       - If > 200 pages, split PDF
       - Copy/rename to avoid Chinese filename issues
       - Upload to MinerU (batch upload API)
       - Poll for results
       - Download and extract content
       - Apply split_into_chunks_v2

    Args:
        directory: Path to directory with PDF files or a specific PDF file

    Returns:
        List of chunked document objects with metadata
    """

    # Get MinerU API token
    token = os.getenv("MinerU_API_KEY")
    if not token:
        print("Error: MinerU_API_KEY not found in .env file")
        return []

    # Get PDF file list
    pdf_docs = DocumentLoader(directory, 'pdf').load_documents()
    if not pdf_docs:
        return []

    processed = []
    temp_dir = project_root / "docs" / "temp_pdf"
    temp_dir.mkdir(exist_ok=True)
    results_dir = temp_dir / "results"
    results_dir.mkdir(exist_ok=True)

    headers = {
        "Content-Type": "application/json",
        "Authorization": f"Bearer {token}"
    }

    for doc in pdf_docs:
        pdf_path = Path(doc['source'])
        print(f"\n{'='*60}")
        print(f"Processing: {pdf_path.name}")
        print(f"{'='*60}")

        try:
            # Get page count
            doc_fitz = fitz.open(str(pdf_path))
            page_count = doc_fitz.page_count
            doc_fitz.close()
            print(f"Page count: {page_count}")

            # Calculate number of splits needed
            num_splits = math.ceil(page_count / 200)
            print(f"Will split into {num_splits} part(s)")

            # Process each split
            all_content = ""
            split_tasks = []

            for split_idx in range(num_splits):
                # Split PDF if needed
                if num_splits > 1:
                    start_page = split_idx * 200
                    end_page = min((split_idx + 1) * 200, page_count)

                    # Create split PDF
                    doc_fitz = fitz.open(str(pdf_path))
                    new_doc = fitz.open()

                    for page_idx in range(start_page, end_page):
                        new_doc.insert_pdf(doc_fitz, from_page=page_idx, to_page=page_idx)

                    # Save split PDF with ASCII-only filename
                    split_filename = f"split_{split_idx}_{int(time.time())}.pdf"
                    split_path = temp_dir / split_filename
                    new_doc.save(str(split_path))
                    new_doc.close()
                    doc_fitz.close()
                    print(f"  Created split {split_idx + 1}: {split_filename} (pages {start_page + 1}-{end_page})")
                else:
                    # Use original PDF, copy to temp dir with ASCII name
                    split_filename = f"full_{int(time.time())}.pdf"
                    split_path = temp_dir / split_filename
                    shutil.copy(str(pdf_path), str(split_path))
                    print(f"  Copied PDF: {split_filename}")

                # Step 1: Get upload URL from MinerU
                print(f"  Requesting upload URL...")
                upload_url_api = "https://mineru.net/api/v4/file-urls/batch"
                upload_data = {
                    "files": [{"name": split_filename}],
                    "model_version": "vlm",
                    "is_ocr": True,
                    "extra_formats": ["docx"]
                }

                try:
                    upload_response = requests.post(upload_url_api, headers=headers, json=upload_data, timeout=30)

                    if upload_response.status_code != 200:
                        print(f"  Error getting upload URL: {upload_response.status_code}")
                        print(f"  Response: {upload_response.text}")
                        continue

                    upload_result = upload_response.json()
                    if upload_result.get("code") != 0:
                        print(f"  API Error: {upload_result.get('msg')}")
                        continue

                    batch_id = upload_result["data"]["batch_id"]
                    file_url = upload_result["data"]["file_urls"][0]
                    print(f"  Got upload URL, batch_id: {batch_id}")

                    # Step 2: Upload the file
                    print(f"  Uploading file...")
                    with open(split_path, 'rb') as f:
                        upload_result = requests.put(file_url, data=f, timeout=120)

                    if upload_result.status_code == 200:
                        print(f"  Upload successful!")
                        split_tasks.append({
                            "batch_id": batch_id,
                            "split_idx": split_idx,
                            "split_filename": split_filename
                        })
                    else:
                        print(f"  Upload failed: {upload_result.status_code}")

                except Exception as e:
                    print(f"  Upload error: {str(e)}")
                    continue

            # Step 3: Poll for results
            print(f"\n  Waiting for processing to complete...")
            for task in split_tasks:
                batch_id = task["batch_id"]
                split_idx = task["split_idx"]

                max_attempts = 60  # Poll for up to 5 minutes (60 * 5s)
                for attempt in range(max_attempts):
                    try:
                        result_url = f"https://mineru.net/api/v4/extract-results/batch/{batch_id}"
                        result_response = requests.get(result_url, headers=headers, timeout=30)

                        if result_response.status_code == 200:
                            result_data = result_response.json()
                            if result_data.get("code") == 0:
                                # extract_result is a list in batch API
                                extract_results = result_data["data"]["extract_result"]
                                if not extract_results:
                                    print(f"  Split {split_idx + 1}: No results found")
                                    break

                                # Get the first (and only) result
                                extract_result = extract_results[0]
                                state = extract_result.get("state")

                                if state == "done":
                                    # Download the result
                                    zip_url = extract_result.get("full_zip_url")
                                    if zip_url:
                                        print(f"  Split {split_idx + 1} done! Downloading results...")
                                        zip_response = requests.get(zip_url, timeout=120)

                                        if zip_response.status_code == 200:
                                            # Extract content from docx file in zip
                                            with zipfile.ZipFile(io.BytesIO(zip_response.content)) as zf:
                                                file_list = zf.namelist()
                                                # print(f"    Files in zip: {file_list}")

                                                # 优先查找 .docx 文件
                                                docx_files = [f for f in file_list if f.endswith('.docx')]

                                                if docx_files:
                                                    # 只读取第一个 .docx 文件
                                                    docx_name = docx_files[0]
                                                    print(f"    Extracting docx: {docx_name}")
                                                    content = read_docx_from_zip(zf, docx_name)
                                                    all_content += content + "\n\n"
                                                    print(f"    Extracted: {docx_name} ({len(content)} chars)")
                                                else:
                                                    # 回退到 .md 文件
                                                    md_files = [f for f in file_list if f.endswith('.md')]
                                                    if md_files:
                                                        md_name = md_files[0]
                                                        print(f"    No docx found, using md: {md_name}")
                                                        content = zf.read(md_name).decode('utf-8')
                                                        all_content += content + "\n\n"
                                                        print(f"    Extracted: {md_name} ({len(content)} chars)")
                                                    else:
                                                        print("    Warning: No .docx or .md files found in zip!")
                                        break
                                elif state == "failed":
                                    print(f"  Split {split_idx + 1} failed: {extract_result.get('err_msg')}")
                                    break
                                elif state in ["pending", "running", "waiting-file", "converting"]:
                                    progress = extract_result.get("extract_progress", {})
                                    if progress:
                                        print(f"  Split {split_idx + 1}: {state} ({progress.get('extracted_pages', 0)}/{progress.get('total_pages', '?')} pages)")
                                    else:
                                        print(f"  Split {split_idx + 1}: {state}...")
                                    time.sleep(5)
                                else:
                                    print(f"  Split {split_idx + 1}: Unknown state '{state}'")
                                    time.sleep(5)
                        else:
                            print(f"  Error checking result: {result_response.status_code}")
                            time.sleep(5)

                    except Exception as e:
                        print(f"  Polling error: {str(e)}")
                        time.sleep(5)

                print(f"  Split {split_idx + 1} processing completed")

            # Step 4: Process content with split_into_chunks_v2
            if all_content.strip():
                print(f"\n  Chunking content ({len(all_content)} chars)...")
                chunks = split_into_chunks_v2(all_content)

                for i, chunk in enumerate(chunks):
                    # Handle both string and Document types
                    if hasattr(chunk, 'page_content'):
                        content = chunk.page_content
                    else:
                        content = str(chunk)

                    processed.append({
                        'doc_id': f"{doc['id']}_chunk_{i}",
                        'content': content,
                        'source': doc['source'],
                        'original_id': doc['id'],
                        'chunk_index': i,
                    })

                print(f"  Created {len(chunks)} chunks")
            else:
                print(f"  Warning: No content extracted from PDF")

        except Exception as e:
            print(f"Error processing {pdf_path.name}: {str(e)}")
            import traceback
            traceback.print_exc()
            continue

    # Cleanup temp directory
    try:
        shutil.rmtree(temp_dir)
        print(f"\nCleaned up temp directory")
    except:
        pass

    print(f"\n{'='*60}")
    print(f"Total: {len(processed)} chunks from {len(pdf_docs)} PDF documents")
    print(f"{'='*60}\n")
    return processed

# Example usage
if __name__ == "__main__":
    # processed_docs = process_documents_v2("./docs/real_docs", "txt") # test.txt is a sample file for testing
    # processed_docs = process_documents_v2("./docs/test.txt") # test.txt is a sample file for testing
    processed_docs = process_documents_pdf("./docs/pdf_docs/上海芯导电子科技股份有限公司财报_2025.pdf")
    print(f"Total chunks created: {len(processed_docs)}")
    # print(processed_docs)  # Print first chunk for verification